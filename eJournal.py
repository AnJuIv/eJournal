import sys
import sqlite3
import os
from openpyxl import Workbook
from docx import Document
from PyQt6 import uic
from PyQt6.QtWidgets import QApplication, QMainWindow, QWidget, QStackedWidget, QMessageBox, QTableWidgetItem


def log_operation(operation, user_id):
    # Функция для записи логов в базу данных.
    conn = sqlite3.connect(db_path)
    cursor = conn.cursor()
    cursor.execute(
        "INSERT INTO Logs (Operation, ID_User) VALUES (?, ?)",
        (operation, user_id))
    conn.commit()
    conn.close()


def path_to_res(relative_path):
    if getattr(sys, 'frozen', False):
        base_path = sys._MEIPASS
    else:
        base_path = os.path.abspath(".")
    return os.path.join(base_path, relative_path)


db_path = os.path.join(
    os.environ["USERPROFILE"], "Документы\eJournal\scheduleDB.sqlite")  # сделать обработчик ошибок и прописать вариант с Documents


class ElZhur(QMainWindow):
    def __init__(self):
        super().__init__()
        uic.loadUi(path_to_res("elzhur.ui"), self)
        self.user = []
        self.aut.clicked.connect(self.Authorization)
        self.acc_stud.clicked.connect(self.studAc)
        self.back_stud.clicked.connect(self.Back)
        self.back.clicked.connect(self.Back_Admin)
        self.save_changes_stud.clicked.connect(self.Save_email)
        self.save_changes.clicked.connect(self.Save_admin)
        self.acc_adm.clicked.connect(self.adminAc)
        self.show_schedule.clicked.connect(self.Show_Schedule)
        self.to_export.clicked.connect(self.To_Export)
        self.to_export_2.clicked.connect(self.To_Export)
        self.back_to_shedule.clicked.connect(self.To_Schedule)
        self.create_report.clicked.connect(self.Create_Report)
        self.export_report.clicked.connect(self.Export_report)
        self.export_schedule.clicked.connect(self.Export_schedule)
        self.out.clicked.connect(self.Logout)
        self.out_admin.clicked.connect(self.Logout)
        self.to_add_user.clicked.connect(self.To_Add_User)
        self.back_2.clicked.connect(self.To_Schedule)
        self.add_user.clicked.connect(self.Add_User)

    def Add_User(self):
        name = self.name_2.text()
        surname = self.surname_2.text()
        patronymic = self.patronymic_2.text()
        email = self.email_2.text()
        login = self.login_2.text()
        password = self.password_2.text()
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        try:
            if (len(email.split("@")) != 2 or email.split(".")[-1] not in ['ru', 'com', 'net']):
                QMessageBox.warning(self, "Некорректный email",
                                    f"Неправильно задан email")
            else:
                cursor.execute("SELECT id FROM Groups WHERE group_name=?",
                               (self.groups_to_add.currentText(),))
                group = cursor.fetchone()
                print(group[0])
                cursor.execute("INSERT INTO Student(login, password, name, surname, patronymic, email, group_id) VALUES (?, ?, ?, ?, ?, ?, ?)",
                               (login, password, name, surname, patronymic, email, group[0]))
                conn.commit()
                QMessageBox.information(
                    self, "Успех", "Данные успешно сохранены.")
                self.name_2.setText('')
                self.surname_2.setText('')
                self.patronymic_2.setText('')
                self.email_2.setText('')
                self.login_2.setText('')
                self.password_2.setText('')
        except Exception as e:
            QMessageBox.warning(self, "Ошибка сохранения",
                                f"Не удалось сохранить данные: {str(e)}")

    def To_Add_User(self):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT group_name FROM Groups")
        group_names = cursor.fetchall()
        for i in group_names:
            self.groups_to_add.addItem(*i)
        conn.close()
        self.stackedWidget.setCurrentIndex(6)

    def Logout(self):
        self.stackedWidget.setCurrentIndex(0)
        self.report.setPlainText('')
        self.user = []

    def Export_schedule(self):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        schedule_data = []

        if self.user[7] == 3:
            for row in range(6):
                row_data = []
                for column in range(6):
                    item = self.tableWidget_2.item(row, column)
                    item = item.text().replace('\n', ' ')
                    row_data.append(item if item else "")
                schedule_data.append(row_data)
        else:
            for row in range(6):
                row_data = []
                for column in range(6):
                    item = self.tableWidget.item(row, column)
                    item = item.text().replace('\n', ' ')
                    row_data.append(item if item else "")
                schedule_data.append(row_data)

        workbook = Workbook()
        sheet = workbook.active
        sheet.title = "Расписание"

        headers = ['Понедельник', 'Вторник', 'Среда',
                   'Четверг', 'Пятница', 'Суббота']
        sheet.append(headers)

        for row in schedule_data:
            sheet.append(row)
        workbook.save(os.path.join(
            os.environ["USERPROFILE"], "Документы\schedule_export.xlsx"))  # сделать обработчик ошибок и прописать вариант с Documents

        QMessageBox.information(
            self, "Экспорт", "Расписание успешно экспортировано в schedule_export.xlsx")
        conn.close()
        log_operation("Экспорт расписания", self.user[0])

    def Export_report(self):
        doc = Document()
        doc.add_heading('Отчет', level=1)
        doc.add_paragraph(self.report.toPlainText())
        doc.save(os.path.join(
            os.environ["USERPROFILE"], "Документы/report.docx"))
        QMessageBox.information(
            self, "Экспорт", "Отчет успешно экспортирован в report.docx")
        log_operation("Экспорт отчета", self.user[0])

    def Create_Report(self):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        if self.user[7] == 3:
            cursor.execute("SELECT id FROM Groups WHERE group_name=?",
                           (self.groups.currentText(),))
            group = cursor.fetchone()[0]
        else:
            group = self.user[7]
        subj = self.subjects.currentText()
        if subj == "Все предметы":
            cursor.execute("""
                SELECT COUNT(*)
                FROM Schedule
                WHERE group_id=?
            """, (group,))
            count = cursor.fetchone()[0]
        else:
            cursor.execute("""
                SELECT COUNT(*)
                FROM Schedule
                WHERE group_id=? AND subject=?
            """, (group, subj))
            count = cursor.fetchone()[0]
        period = self.period.currentText()
        self.report.setPlainText(self.report.toPlainText(
        ) + f"По предмету {subj} в {period} отведено (в часах): {count * 1.5 * {'Год': 52, 'Месяц': 4,'Неделя':1}[period]}.\n")
        log_operation("Создание отчета", self.user[0])
        conn.close()

    def To_Schedule(self):
        if self.user[7] == 3:
            self.stackedWidget.setCurrentIndex(4)
        else:
            self.stackedWidget.setCurrentIndex(3)


    def Back_Admin(self):
        self.stackedWidget.setCurrentIndex(4)

    def Back(self):
        self.stackedWidget.setCurrentIndex(3)

    def Authorization(self)    def To_Export(self):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        if self.user[7] == 3:
            cursor.execute("SELECT id FROM Groups WHERE group_name=?",
                           (self.groups.currentText(),))
            group = cursor.fetchone()
            cursor.execute(
                "SELECT subject FROM Schedule WHERE group_id=? GROUP BY subject", group)
            subjects = cursor.fetchall()
        else:
            group = self.user[7]
            cursor.execute(
                "SELECT subject FROM Schedule WHERE group_id=? GROUP BY subject", (group,))
            subjects = cursor.fetchall()
        for i in subjects:
            self.subjects.addItem(*i)
        self.stackedWidget.setCurrentIndex(5)
        conn.close()

    def Show_Schedule(self):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        cursor.execute("SELECT id FROM Groups WHERE group_name=?",
                       (self.groups.currentText(),))
        id = cursor.fetchone()
        self.populate_schedule(*id)

    def Save_email(self):
        email = self.email.text()
        student_id = self.user[0]
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        try:
            cursor.execute(
                "UPDATE Student SET email=? WHERE id=?", (email, student_id))
            conn.commit()
            QMessageBox.information(
                self, "Успех", "Email успешно обновлен.")
            cursor.execute("SELECT * FROM Student WHERE id=?", (student_id,))
            self.user = cursor.fetchone()
            self.Back()
        except Exception as e:
            QMessageBox.warning(self, "Ошибка сохранения",
                                f"Не удалось обновить данные: {str(e)}")
        finally:
            conn.close()
:
        log = self.login.text()
        pas = self.password.text()
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()
        #Поиск пользователя в базе данных
        cursor.execute(
            "SELECT * FROM Student WHERE login=? AND password=?", (log, pas))
        user = cursor.fetchone()
        # Успешно найден, авторизация
        if user:
            self.user = user
            group_id = user[7]
            if group_id == 3:
                cursor.execute(
                    "SELECT group_name FROM Groups WHERE group_name IS NOT 'Admin'")
                group_names = cursor.fetchall()
                for i in group_names:
                    self.groups.addItem(*i)
                self.stackedWidget.setCurrentIndex(4)
            else:
                cursor.execute(
                    "SELECT group_name FROM Groups WHERE id=?", (group_id,))
                group_name = cursor.fetchone()
                if group_name:
                    self.group_name.setText(group_name[0])

                self.populate_schedule(group_id)

                self.stackedWidget.setCurrentIndex(3)
            log_operation("Авторизация", user[0])
        #ошибка входа
        else:
            QMessageBox.warning(self, "Ошибка входа",
                                "Неверный логин или пароль.")

        conn.close()

    def studAc(self):
        self.email.setText(self.user[3])
        self.name.setText(self.user[4])
        self.surname.setText(self.user[5])
        self.patronymic.setText(self.user[6])
        self.stackedWidget.setCurrentIndex(1)

    def adminAc(self):
        self.email_admin.setText(self.user[3])
        self.name_admin.setText(self.user[4])
        self.surname_admin.setText(self.user[5])
        self.patronymic_admin.setText(self.user[6])
        self.stackedWidget.setCurrentIndex(2)

    def populate_schedule(self, group_id):
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        try:
            # Получение расписания для заданной группы с информацией о преподавателе
            cursor.execute("""
                SELECT S.subject, S.weekday, C.class_number, L.surname, L.name
                FROM Schedule S
                JOIN Classes C ON S.class_number = C.class_number
                JOIN Lecturer L ON S.lecturer_id = L.id
                WHERE S.group_id = ?
                ORDER BY S.weekday, C.class_number
            """, (group_id,))
            schedule = cursor.fetchall()

            # Получение всех классов
            cursor.execute("SELECT * FROM Classes")
            classes = {i[0]: str(i[1]) + '-' + str(i[2]) for i in cursor.fetchall()}

            # Инициализация словаря расписания
            schedule_dict = {day: [''] * 6 for day in ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']}

            # Заполнение расписания
            for subject, weekday, class_number, lecturer_surname, lecturer_name in schedule:
                day_index = ['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday'].index(weekday)
                # Форматирование строки с предметом и ФИО преподавателя
                lecturer_initials = f"{lecturer_surname} {lecturer_name[0]}."  # Фамилия и первая буква имени
                schedule_dict[weekday][class_number - 1] = f"{subject}\n{lecturer_initials}"

            # Определение таблицы для отображения
            if self.user[7] == 3:
                table_widget = self.tableWidget_2
            else:
                table_widget = self.tableWidget

            # Настройка таблицы
            table_widget.clear()
            table_widget.setColumnCount(6)
            table_widget.setRowCount(6)
            table_widget.setHorizontalHeaderLabels(['Понедельник', 'Вторник', 'Среда', 'Четверг', 'Пятница', 'Суббота'])

            # Заполнение таблицы данными расписания
            for day_index, day in enumerate(['Monday', 'Tuesday', 'Wednesday', 'Thursday', 'Friday', 'Saturday']):
                for class_index in range(6):
                    subject_display = classes.get(class_index + 1, '') + '\n' + schedule_dict[day][class_index]
                    if subject_display.strip():  # Проверка на пустую строку
                        table_widget.setItem(class_index, day_index, QTableWidgetItem(subject_display))

            # Настройка размеров столбцов и строк
            table_widget.resizeColumnsToContents()
            for row in range(6):
                table_widget.setRowHeight(row, 50)

        except Exception as e:
            QMessageBox.warning(self, "Ошибка", f"Не удалось загрузить расписание: {str(e)}")

        finally:
            conn.close()

    def Save_admin(self):
        name = self.name_admin.text()
        surname = self.surname_admin.text()
        patronymic = self.patronymic_admin.text()
        email = self.email_admin.text()
        admin_id = self.user[0]
        conn = sqlite3.connect(db_path)
        cursor = conn.cursor()

        try:
            #Обновление данных
            cursor.execute("UPDATE Student SET name=?, surname=?, patronymic=?, email=? WHERE id=?",
                           (name, surname, patronymic, email, admin_id))
            conn.commit()
            QMessageBox.information(
                self, "Успех", "Данные успешно обновлены.")
            cursor.execute("SELECT * FROM Student WHERE id=?", (admin_id,))
            self.user = cursor.fetchone()
            self.Back_Admin()
        except Exception as e:
            QMessageBox.warning(self, "Ошибка сохранения",
                                f"Не удалось обновить данные: {str(e)}")
        finally:
            conn.close()


app = QApplication(sys.argv)
window = ElZhur()
#Название окна
window.setWindowTitle("Электронный журнал")
try:
    window.show()
except Exception as e:  # Обрабатываем исключения при запуске
    QMessageBox.warning(window, "Ошибка", f"Не удалось запустить приложение: {str(e)}")
app.exec()